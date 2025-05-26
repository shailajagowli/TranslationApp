import io
from fastapi import FastAPI, UploadFile, File, Response
from fastapi.middleware.cors import CORSMiddleware
from hashlib import sha256
from difflib import SequenceMatcher
from time import time

from openai import OpenAI
from docx import Document
import openpyxl

from sqlalchemy import create_engine, Column, String
from sqlalchemy.orm import sessionmaker, declarative_base

# --- Setup OpenAI client ---
API_KEY = ""
client = OpenAI(api_key=API_KEY)

ENERGY_DOMAIN_PROMPT = """You are a professional translator specialized in the energy domain.
Translate the following English text into German with technical accuracy, consistent terminology, and proper grammar.
Ensure consistent formatting. Only output the German translation."""

# --- Database setup ---
Base = declarative_base()
engine = create_engine("sqlite:///translations.db")
Session = sessionmaker(bind=engine)

class Translation(Base):
    __tablename__ = "translations"
    hash = Column(String, primary_key=True)
    original = Column(String)
    translated = Column(String)

Base.metadata.create_all(engine)

app = FastAPI()

# CORS middleware for frontend interaction
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # your React dev server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=[
        "X-Efficiency",
        "X-Memory-Hits",
        "X-Fuzzy-Hits",
        "X-OpenAI-Hits",
        "X-Time-Seconds",
        "Content-Disposition",
        "Content-Type",
    ],
)

# --- Utility functions ---

def get_hash(text: str) -> str:
    return sha256(text.encode("utf-8")).hexdigest()

# DOCX extraction & replacement
def extract_texts_docx(doc_bytes):
    doc = Document(io.BytesIO(doc_bytes))
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text)
    return texts

def replace_texts_docx(original_bytes, translated_texts):
    doc = Document(io.BytesIO(original_bytes))
    idx = 0
    for para in doc.paragraphs:
        if para.text.strip() and idx < len(translated_texts):
            para.clear()  # Clear runs but keep style
            para.add_run(translated_texts[idx])
            idx += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() and idx < len(translated_texts):
                    cell.text = translated_texts[idx]
                    idx += 1
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# XLSX extraction & replacement
def extract_texts_xlsx(file_bytes):
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes))
    texts = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str) and cell.value.strip():
                    texts.append(cell.value)
    return texts

def replace_texts_xlsx(original_bytes, translated_texts):
    wb = openpyxl.load_workbook(io.BytesIO(original_bytes))
    idx = 0
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str) and cell.value.strip() and idx < len(translated_texts):
                    cell.value = translated_texts[idx]
                    idx += 1
    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()

# TXT extraction & replacement
def extract_texts_txt(file_bytes):
    text = file_bytes.decode("utf-8")
    return [line for line in text.splitlines() if line.strip()]

def replace_texts_txt(original_bytes, translated_texts):
    return ("\n".join(translated_texts)).encode("utf-8")

# Fuzzy matching DB lookup
def find_fuzzy_match(session, text, threshold=0.9):
    all_translations = session.query(Translation).all()
    best_match = None
    best_ratio = 0
    for item in all_translations:
        ratio = SequenceMatcher(None, text, item.original).ratio()
        if ratio > best_ratio and ratio >= threshold:
            best_match = item
            best_ratio = ratio
    return (best_match, best_ratio) if best_match else (None, 0)

# Core batch translation logic
def translate_batch(texts: list[str]) -> dict:
    session = Session()
    results = []
    memory_hits = 0
    fuzzy_hits = 0
    openai_hits = 0
    start_time = time()

    for text in texts:
        h = get_hash(text)
        existing = session.query(Translation).filter_by(hash=h).first()
        if existing:
            memory_hits += 1
            results.append({
                "original": text,
                "translated": existing.translated,
                "match_type": "exact",
                "similarity": 1.0,
            })
            continue

        fuzzy_match, ratio = find_fuzzy_match(session, text)
        if fuzzy_match and ratio > 0.9:
            fuzzy_hits += 1
            results.append({
                "original": text,
                "translated": fuzzy_match.translated,
                "match_type": "fuzzy",
                "similarity": ratio,
            })
            continue

        # OpenAI call
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": ENERGY_DOMAIN_PROMPT},
                {"role": "user", "content": text},
            ],
            temperature=0.2,
        )
        translated = response.choices[0].message.content
        session.add(Translation(hash=h, original=text, translated=translated))
        session.commit()
        openai_hits += 1
        results.append({
            "original": text,
            "translated": translated,
            "match_type": "new",
            "similarity": None,
        })

    duration = time() - start_time
    session.close()

    return {
        "translations": results,
        "stats": {
            "total": len(texts),
            "memory_hits": memory_hits,
            "fuzzy_hits": fuzzy_hits,
            "openai_hits": openai_hits,
            "efficiency": round((memory_hits + fuzzy_hits) / len(texts) * 100, 2) if texts else 0,
            "time_seconds": round(duration, 2),
        },
    }



from fastapi.responses import JSONResponse

@app.post("/translate")
async def translate(file: UploadFile = File(...)):
    content = await file.read()
    ext = file.filename.split(".")[-1].lower()

    # Extract texts
    if ext == "docx":
        texts = extract_texts_docx(content)
    elif ext == "xlsx":
        texts = extract_texts_xlsx(content)
    elif ext == "txt":
        texts = extract_texts_txt(content)
    else:
        return Response(content="Unsupported file type", status_code=400)

    result = translate_batch(texts)
    # Return full translation objects + stats for review
    return JSONResponse(content=result)

from pydantic import BaseModel

class TranslationItem(BaseModel):
    original: str
    translated: str
    match_type: str
    similarity: float | None

class GenerateRequest(BaseModel):
    translations: list[TranslationItem]
    filename: str
    file_bytes: str  # base64 encoded original file bytes
    ext: str

import base64

@app.post("/generate")
async def generate(request: GenerateRequest):
    file_bytes = base64.b64decode(request.file_bytes)
    ext = request.ext
    translated_texts = [item.translated for item in request.translations]

    if ext == "docx":
        translated_file_bytes = replace_texts_docx(file_bytes, translated_texts)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif ext == "xlsx":
        translated_file_bytes = replace_texts_xlsx(file_bytes, translated_texts)
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif ext == "txt":
        translated_file_bytes = replace_texts_txt(file_bytes, translated_texts)
        media_type = "text/plain"
    else:
        return Response(content="Unsupported file type", status_code=400)

    return Response(
        content=translated_file_bytes,
        media_type=media_type,
        headers={
            "Content-Disposition": f'attachment; filename="translated_{request.filename}"',
        }
    )
